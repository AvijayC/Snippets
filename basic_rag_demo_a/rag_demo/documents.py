from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


SUPPORTED_SUFFIXES = {".md", ".markdown", ".txt", ".html", ".htm", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    id: str
    path: Path
    title: str
    text: str
    metadata: dict[str, str]


@dataclass
class DocumentChunk:
    id: str
    document_id: str
    source_path: str
    title: str
    text: str
    chunk_index: int
    metadata: dict[str, str]


def iter_supported_files(path: Path) -> Iterable[Path]:
    if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES:
        yield path
        return
    if not path.exists():
        return
    for file_path in sorted(path.rglob("*")):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_SUFFIXES:
            yield file_path


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix in {".html", ".htm"}:
        text = _extract_html(path)
    elif suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    else:
        raise ValueError(f"Unsupported document type: {path}")
    text = normalize_text(text)
    title = infer_title(path, text)
    doc_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    return LoadedDocument(
        id=doc_hash[:24],
        path=path,
        title=title,
        text=text,
        metadata={"suffix": suffix, "sha256": doc_hash},
    )


def chunk_document(document: LoadedDocument, chunk_chars: int, overlap_chars: int) -> list[DocumentChunk]:
    if not document.text:
        return []
    chunk_chars = max(500, chunk_chars)
    overlap_chars = max(0, min(overlap_chars, chunk_chars // 2))
    chunks = []
    start = 0
    index = 0
    text = document.text
    while start < len(text):
        end = min(len(text), start + chunk_chars)
        if end < len(text):
            boundary = max(text.rfind("\n", start, end), text.rfind(". ", start, end))
            if boundary > start + chunk_chars // 2:
                end = boundary + 1
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append(
                DocumentChunk(
                    id=f"{document.id}:{index}",
                    document_id=document.id,
                    source_path=str(document.path),
                    title=document.title,
                    text=chunk_text,
                    chunk_index=index,
                    metadata=document.metadata | {"chunk_index": str(index)},
                )
            )
            index += 1
        if end >= len(text):
            break
        start = max(0, end - overlap_chars)
    return chunks


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_title(path: Path, text: str) -> str:
    for line in text.splitlines():
        cleaned = line.strip().lstrip("#").strip()
        if cleaned:
            return cleaned[:120]
    return path.stem.replace("_", " ").replace("-", " ").title()


def _extract_html(path: Path) -> str:
    html = path.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()
        return soup.get_text("\n")
    except Exception:
        return re.sub(r"<[^>]+>", " ", html)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("PDF support requires the pypdf package.") from exc
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except Exception as exc:
        raise RuntimeError("DOCX support requires the python-docx package.") from exc
    document = docx.Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)
